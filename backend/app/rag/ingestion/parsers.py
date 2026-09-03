import hashlib
import mimetypes
import uuid
import zipfile

from pathlib import Path

import pymupdf

from docx import Document as DocxDocument
from PIL import Image

from app.rag.ingestion.types import (
    ParsedPage,
    ParsedVisual,
    ParseResult,
)


def _checksum(
    data: bytes,
) -> str:

    return hashlib.sha256(
        data
    ).hexdigest()


def _relative_path(
    root: Path,
    path: Path,
) -> str:

    return (
        path
        .relative_to(root)
        .as_posix()
    )


def _image_dimensions(
    path: Path,
) -> tuple[int | None, int | None]:

    try:
        with Image.open(path) as image:
            return image.size

    except Exception:
        return None, None


def parse_pdf(
    file_path: Path,
    document_id: uuid.UUID,
    storage_root: Path,
) -> ParseResult:

    pages: list[ParsedPage] = []

    visuals: list[ParsedVisual] = []

    page_output = (
        storage_root
        / "page_images"
        / str(document_id)
    )

    image_output = (
        storage_root
        / "extracted_images"
        / str(document_id)
    )

    page_output.mkdir(
        parents=True,
        exist_ok=True,
    )

    image_output.mkdir(
        parents=True,
        exist_ok=True,
    )

    asset_index = 0

    pdf = pymupdf.open(
        file_path
    )

    try:

        for page_idx in range(
            len(pdf)
        ):
            page = pdf[page_idx]

            page_number = (
                page_idx + 1
            )

            text = page.get_text(
                "text",
                sort=True,
            )

            pages.append(
                ParsedPage(
                    page_number=page_number,
                    text=text,
                )
            )

            # Render the entire page.
            page_image_path = (
                page_output
                / f"page_{page_number:04d}.png"
            )

            pixmap = page.get_pixmap(
                dpi=144,
                alpha=False,
            )

            pixmap.save(
                str(page_image_path)
            )

            page_bytes = (
                page_image_path
                .read_bytes()
            )

            visuals.append(
                ParsedVisual(
                    asset_index=asset_index,
                    page_number=page_number,
                    asset_type="page",
                    storage_path=_relative_path(
                        storage_root,
                        page_image_path,
                    ),
                    mime_type="image/png",
                    checksum_sha256=_checksum(
                        page_bytes
                    ),
                    width_px=pixmap.width,
                    height_px=pixmap.height,
                )
            )

            asset_index += 1

            # Extract embedded raster images.
            for image_number, image_info in enumerate(
                page.get_images(
                    full=True
                ),
                start=1,
            ):
                xref = image_info[0]

                extracted = (
                    pdf.extract_image(
                        xref
                    )
                )

                image_bytes = (
                    extracted["image"]
                )

                extension = (
                    extracted.get(
                        "ext",
                        "png",
                    )
                )

                output_path = (
                    image_output
                    / (
                        f"page_{page_number:04d}"
                        f"_image_{image_number:03d}"
                        f".{extension}"
                    )
                )

                output_path.write_bytes(
                    image_bytes
                )

                mime_type = (
                    mimetypes.guess_type(
                        output_path.name
                    )[0]
                )

                visuals.append(
                    ParsedVisual(
                        asset_index=asset_index,
                        page_number=page_number,
                        asset_type="embedded_image",
                        storage_path=_relative_path(
                            storage_root,
                            output_path,
                        ),
                        mime_type=mime_type,
                        checksum_sha256=_checksum(
                            image_bytes
                        ),
                        width_px=extracted.get(
                            "width"
                        ),
                        height_px=extracted.get(
                            "height"
                        ),
                    )
                )

                asset_index += 1

    finally:
        pdf.close()

    return ParseResult(
        pages=pages,
        visuals=visuals,
        page_count=len(pages),
    )


def parse_docx(
    file_path: Path,
    document_id: uuid.UUID,
    storage_root: Path,
) -> ParseResult:

    document = DocxDocument(
        file_path
    )

    text_parts: list[str] = []

    for paragraph in document.paragraphs:

        value = (
            paragraph.text.strip()
        )

        if value:
            text_parts.append(
                value
            )

    for table in document.tables:

        for row in table.rows:

            cells = [
                cell.text.strip()
                for cell in row.cells
            ]

            text_parts.append(
                " | ".join(cells)
            )

    text = "\n".join(
        text_parts
    )

    visuals: list[ParsedVisual] = []

    image_output = (
        storage_root
        / "extracted_images"
        / str(document_id)
    )

    image_output.mkdir(
        parents=True,
        exist_ok=True,
    )

    asset_index = 0

    with zipfile.ZipFile(
        file_path
    ) as archive:

        media_files = [
            name
            for name in archive.namelist()
            if name.startswith(
                "word/media/"
            )
        ]

        for media_name in media_files:

            image_bytes = (
                archive.read(
                    media_name
                )
            )

            extension = (
                Path(media_name)
                .suffix
                .lower()
            )

            if not extension:
                extension = ".bin"

            output_path = (
                image_output
                / (
                    f"image_{asset_index:04d}"
                    f"{extension}"
                )
            )

            output_path.write_bytes(
                image_bytes
            )

            width, height = (
                _image_dimensions(
                    output_path
                )
            )

            visuals.append(
                ParsedVisual(
                    asset_index=asset_index,
                    page_number=None,
                    asset_type="embedded_image",
                    storage_path=_relative_path(
                        storage_root,
                        output_path,
                    ),
                    mime_type=(
                        mimetypes.guess_type(
                            output_path.name
                        )[0]
                    ),
                    checksum_sha256=_checksum(
                        image_bytes
                    ),
                    width_px=width,
                    height_px=height,
                )
            )

            asset_index += 1

    return ParseResult(
        pages=[
            ParsedPage(
                page_number=None,
                text=text,
            )
        ],
        visuals=visuals,
        page_count=None,
    )


def parse_text_file(
    file_path: Path,
) -> ParseResult:

    text = file_path.read_text(
        encoding="utf-8-sig",
        errors="replace",
    )

    return ParseResult(
        pages=[
            ParsedPage(
                page_number=None,
                text=text,
            )
        ],
        visuals=[],
        page_count=None,
    )


def parse_document(
    file_path: Path,
    document_id: uuid.UUID,
    storage_root: Path,
) -> ParseResult:

    extension = (
        file_path
        .suffix
        .lower()
    )

    if extension == ".pdf":

        return parse_pdf(
            file_path,
            document_id,
            storage_root,
        )

    if extension == ".docx":

        return parse_docx(
            file_path,
            document_id,
            storage_root,
        )

    if extension in {
        ".txt",
        ".md",
    }:

        return parse_text_file(
            file_path
        )

    raise ValueError(
        f"No parser available for {extension}"
    )